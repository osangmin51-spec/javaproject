from pathlib import Path
import os

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path.cwd()
OUT = ROOT / "deliverables"
OUT.mkdir(exist_ok=True)
PREVIEW = OUT / "mock-stock-website-demo-preview.png"
VIDEO = OUT / "mock-stock-website-demo.webm"
UI_SCREENSHOTS = [
    (OUT / "ui-screenshots" / "ui-crop-01-market.png", "화면 1. 시장 시세: 거래량 상위 종목, 검색창, 즐겨찾기, 페이지 이동"),
    (OUT / "ui-screenshots" / "ui-crop-02-detail-order.png", "화면 2. 종목 상세: 현재가, 등락률, 거래량, 회사 설명, 매수/매도 입력"),
    (OUT / "ui-screenshots" / "ui-crop-03-portfolio.png", "화면 3. 보유 탭: 보유 수량, 평가금액, 손익, 수익률, 매도 입력"),
]


doc = Document()
section = doc.sections[0]
section.top_margin = Inches(0.8)
section.bottom_margin = Inches(0.75)
section.left_margin = Inches(0.85)
section.right_margin = Inches(0.85)

styles = doc.styles
styles["Normal"].font.name = "맑은 고딕"
styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "맑은 고딕")
styles["Normal"].font.size = Pt(10.2)

for name, size, color in [
    ("Heading 1", 16, RGBColor(18, 32, 51)),
    ("Heading 2", 13, RGBColor(31, 95, 191)),
    ("Heading 3", 11.5, RGBColor(24, 34, 48)),
]:
    style = styles[name]
    style.font.name = "맑은 고딕"
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "맑은 고딕")
    style.font.size = Pt(size)
    style.font.bold = True
    style.font.color.rgb = color


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text, bold=False, color=None):
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(str(text))
    r.font.name = "맑은 고딕"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "맑은 고딕")
    r.font.size = Pt(9.2)
    r.font.bold = bold
    if color:
        r.font.color.rgb = color
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_table(headers, rows, widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for i, h in enumerate(headers):
        shade_cell(table.rows[0].cells[i], "122033")
        set_cell_text(table.rows[0].cells[i], h, bold=True, color=RGBColor(255, 255, 255))
    for row_index, row in enumerate(rows):
        cells = table.add_row().cells
        for i, value in enumerate(row):
            if row_index % 2 == 1:
                shade_cell(cells[i], "F6F9FD")
            set_cell_text(cells[i], value)
    if widths:
        for row in table.rows:
            for idx, width in enumerate(widths):
                row.cells[idx].width = Inches(width)
    doc.add_paragraph()
    return table


def add_bullets(items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(3)
        p.add_run(item)


def add_numbered(items):
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.paragraph_format.space_after = Pt(3)
        p.add_run(item)


def add_code(text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.15)
    p.paragraph_format.right_indent = Inches(0.15)
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(text)
    r.font.name = "Consolas"
    r.font.size = Pt(8.8)
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), "F1F5F9")
    p._p.get_or_add_pPr().append(shd)


def paragraph(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.add_run(text)


title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = title.add_run("Java 모의주식투자 웹앱 개인 프로젝트 보고서")
r.font.name = "맑은 고딕"
r._element.rPr.rFonts.set(qn("w:eastAsia"), "맑은 고딕")
r.font.size = Pt(22)
r.font.bold = True
r.font.color.rgb = RGBColor(18, 32, 51)

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = subtitle.add_run("한국투자증권 KIS Open API 기반 모의투자 시스템")
r.font.name = "맑은 고딕"
r._element.rPr.rFonts.set(qn("w:eastAsia"), "맑은 고딕")
r.font.size = Pt(12)
r.font.color.rgb = RGBColor(102, 112, 133)

doc.add_paragraph()
add_table(["항목", "내용"], [
    ["프로젝트명", "Java 프로젝트 모의주식"],
    ["주요 기술", "Java HttpServer, HttpClient, Thread, Collection, JDBC, MySQL"],
    ["외부 데이터", "한국투자증권 KIS Open API"],
    ["저장 방식", "MySQL 우선 저장 + TSV fallback"],
    ["실행 주소", "http://localhost:8080/"],
    ["시연 영상", "deliverables/mock-stock-website-demo.webm"],
], widths=[1.7, 5.7])


doc.add_heading("1. 프로젝트 개요", level=1)
paragraph("본 프로젝트는 Java 표준 라이브러리 기반의 모의주식투자 웹앱이다. 사용자는 웹 브라우저에서 바로 국내 주식 종목을 확인하고, 관심 있는 종목을 클릭해 상세 정보를 본 다음 매수할 수 있다. 이미 산 종목은 보유 탭에서 수량을 지정해 매도한다.")
paragraph("핵심 목표는 단순한 주문 연습이 아니라 실제 시세, 거래량, 가격 그래프, 포트폴리오 손익을 한 화면 흐름으로 연결하는 것이다. 한국투자증권 KIS Open API를 이용해 국내주식 현재가와 거래량 정보를 가져오고, 사용자는 종목 상세 화면에서 가격 변화를 확인한 뒤 매수와 매도를 진행한다.")
add_bullets([
    "실제 외부 증권 API를 활용한 Java 웹 프로젝트 구현",
    "사용자 매수/매도 흐름과 포트폴리오 손익 계산 구현",
    "MySQL 우선 저장과 TSV fallback을 통한 모의 계좌/보유/거래 기록 유지",
    "쓰레드 기반 시세 갱신과 소켓 구독 구조 실험",
    "검색, 즐겨찾기, 종목 상세, 가격 그래프 UI 구현",
    "다수 클래스/인터페이스 구조와 실제 실행 가능한 웹 UI 구현",
])

doc.add_heading("1.1 한눈에 보는 핵심 구현", level=2)
paragraph("교수님이 프로젝트를 처음 볼 때 가장 먼저 이해해야 하는 부분은 사용자의 화면 조작이 Java 서버, 외부 시세 API, MySQL 저장소로 어떻게 이어지는지이다. 이 프로젝트는 프론트엔드 프레임워크 없이 Java 서버가 HTML/CSS/JavaScript 화면을 직접 내려주고, 브라우저는 필요한 순간 API를 호출해 상태를 다시 받아오는 구조다.")
add_table(["구분", "구현 내용", "발표에서 강조할 점"], [
    ["사용자 화면", "시장 시세, 검색, 즐겨찾기, 종목 상세, 매수/매도, 보유 탭", "사용자가 실제 투자 앱처럼 종목을 보고 판단한 뒤 주문한다."],
    ["외부 시세", "KIS Open API 현재가와 거래량 조회", "더미 데이터가 아니라 외부 시세 연동을 시도한 프로젝트다."],
    ["서버 로직", "MiniHandler가 요청을 받고 MiniProject가 매매와 손익을 계산", "Java 표준 HttpServer 기반으로 직접 요청 흐름을 구현했다."],
    ["저장 구조", "MySQL members, shares, trade_logs + TSV fallback", "서버를 껐다 켜도 계좌와 거래 기록을 유지한다."],
    ["실시간성", "Thread 기반 갱신과 WebSocket 구독 구조 실험", "개인 PC에서 가능한 범위 안에서 실시간성을 설계했다."],
], widths=[1.4, 3.1, 2.9])


doc.add_heading("2. 제안발표 이후 주제 변화", level=1)
paragraph("초기 제안 단계에서는 기본 Java 기능을 웹 화면으로 옮기고 계좌, 매매 기능을 구현하는 방향이었다. 진행 중 모의주식투자라는 주제에는 내부 샘플 가격보다 실제 주식 시세와 거래량이 더 적합하다고 판단했다.")
add_table(["구분", "제안발표 단계", "최종 구현"], [
    ["중심 주제", "기본 Java 기능 구현", "KIS API 기반 모의주식투자 웹앱"],
    ["가격 데이터", "내부 샘플/모의 가격", "한국투자증권 KIS 현재가"],
    ["종목 수", "소수 종목", "거래량 상위 종목 중심"],
    ["화면", "기본 입력/출력 중심", "검색, 즐겨찾기, 상세 그래프, 포트폴리오"],
    ["저장", "메모리 중심", "MySQL 우선 저장 + TSV fallback"],
    ["목적", "기능 구현 연습", "실제 데이터 기반 투자 흐름 구현"],
], widths=[1.3, 2.6, 3.3])


doc.add_heading("3. 자바 기반 프로그램 설계", level=1)
paragraph("프로젝트는 Spring이나 React 없이 Java 표준 기능을 중심으로 구성했다. HttpServer로 웹 서버를 열고, HttpClient로 외부 API를 호출하며, HTML/CSS/JavaScript는 Java 문자열 템플릿에서 렌더링한다. 저장은 MySQL JDBC 구조를 우선 사용하고, MYSQL_URL과 MYSQL_USER가 없거나 MySQL 연결에 실패하면 LocalFileDatabase가 data/local-database.tsv에 저장하는 fallback 구조로 전환한다. MYSQL_PASSWORD는 로컬 계정 설정에 따라 빈 값일 수 있다.")
add_table(["파일", "역할"], [
    ["app/MiniProjectApp.java", "서버 시작, 포트 설정, KIS/API/DB 초기화"],
    ["MiniHandler.java", "URL별 HTTP 요청 라우팅과 JSON 응답 처리"],
    ["MiniProject.java", "모의 계좌, 매매, 포트폴리오, 시세 상태 관리"],
    ["domain/Member.java, Stock.java", "Member, Stock, Share, TradeLog 등 도메인 모델"],
    ["repository/MySqlDatabase.java", "MySQL 연결, 테이블 생성, 데이터 저장/로드"],
    ["repository/LocalFileDatabase.java", "MySQL 연결 실패 시 TSV 파일 저장/로드"],
    ["external/KisQuotePoller.java", "KIS 토큰 발급, 거래량 순위, 현재가 조회, 시세 갱신"],
    ["domain/StockCategories.java", "업종·테마별 위험 설명 타입"],
    ["external/MockBrokerServer.java", "소켓 기반 시세 구독 흐름 실험"],
    ["view/MiniDashboardPage.java", "HTML/CSS/JavaScript 화면 생성"],
    ["util/Json.java", "JSON 문자열 생성과 요청 body 파싱"],
], widths=[2.1, 5.2])

doc.add_heading("3.1 패키지/클래스 다이어그램 요약", level=2)
add_code("""src/main/java
 ├─ app/MiniProjectApp
 ├─ controller/MiniHandler ── service/MiniProject
 │                            ├─ domain/Member, Stock, Share, TradeLog
 │                            ├─ repository/MySqlDatabase
 │                            └─ external/KisQuotePoller, MockBrokerServer
 ├─ view/MiniDashboardPage
 └─ util/Json""")

doc.add_heading("3.2 클래스/라이브러리 구조 상세", level=2)
paragraph("프로젝트는 src/main/java 아래에 app, controller, service, domain, repository, external, view, util 패키지로 분리했다. 발표에서는 아래 표를 기준으로 요청이 어떤 계층을 거쳐 처리되는지 설명하면 된다.")
add_table(["계층", "주요 클래스", "사용 라이브러리/문법", "역할"], [
    ["실행/서버 계층", "app.MiniProjectApp, controller.MiniHandler", "HttpServer, HttpExchange, Executors", "웹 서버 시작, URL 라우팅, 요청/응답 처리"],
    ["핵심 서비스 계층", "MiniProject", "ConcurrentHashMap, ArrayList, Comparator, LocalDateTime", "매수/매도 처리, 포트폴리오 계산, 시세 상태 관리"],
    ["도메인 모델 계층", "Member, Stock, Share, TradeLog, PricePoint", "Java class, 컬렉션 객체", "모의 계좌, 종목, 보유 수량, 거래 기록, 가격 이력 표현"],
    ["외부 시세 계층", "KisQuoteClient, KisQuotePoller, KisWebSocketQuoteClient", "HttpClient, HttpRequest, WebSocket", "KIS REST 현재가 조회, 거래량 상위 종목 갱신, WebSocket 연동 시도"],
    ["저장소 계층", "MySqlDatabase, DatabaseSnapshot", "JDBC, DriverManager", "MySQL 테이블 생성, 데이터 로드/저장"],
    ["화면/직렬화 계층", "view.MiniDashboardPage, util.Json", "HTML/CSS/JavaScript 문자열, JSON 생성/파싱", "브라우저 화면 렌더링과 API 응답 데이터 구성"],
    ["분류/상속 계층", "CompanyProfile, CompanyProfiles, StockCategoryProfile", "abstract class, interface, 구현 클래스", "회사 설명과 업종 위험 설명을 종목 상세 설명에 반영"],
], widths=[1.35, 2.1, 2.2, 2.3])
paragraph("src/main/java 패키지 구조로 옮겼기 때문에, 발표에서는 실제 파일 위치와 역할을 같이 설명할 수 있다. 다만 MiniProject가 아직 매매, 포트폴리오, 시장 상태를 함께 관리하므로 더 큰 프로젝트라면 service 계층을 더 세분화할 수 있다.")

doc.add_heading("3.3 주요 코드 흐름", level=2)
paragraph("아래 코드는 발표 때 보여주기 위한 핵심 부분만 줄인 것이다. 전체 코드를 모두 설명하기보다 HTTP 요청이 들어오고, 외부 시세를 조회하고, 매수 결과를 저장하는 흐름을 중심으로 정리했다.")
paragraph("예를 들어 사용자가 삼성전기 상세 화면에서 매수 버튼을 누르면 브라우저는 /api/stock/buy로 종목명과 수량을 보낸다. MiniHandler는 이 요청을 MiniProject.buyStock()으로 넘기고, MiniProject는 잔액과 수량을 확인한 뒤 보유 종목과 거래 기록을 갱신한다. 마지막으로 현재 선택된 저장소가 변경된 계좌, 보유 주식, 거래 기록을 저장한다. 기본은 MySQL이고, 연결 실패 시 LocalFileDatabase가 TSV 파일로 저장한다.")
add_table(["순서", "처리 위치", "설명"], [
    ["1", "브라우저", "종목 상세 화면에서 수량 입력 후 매수 버튼 클릭"],
    ["2", "MiniHandler", "POST /api/stock/buy 요청 body를 읽고 MiniProject로 전달"],
    ["3", "MiniProject", "현재가 × 수량 계산, 잔액 부족 여부 확인, 보유 종목 갱신"],
    ["4", "ProjectDatabase", "MySQL 또는 TSV 저장소에 members, shares, trade_logs 데이터 저장"],
    ["5", "브라우저", "GET /api/state로 최신 현금, 평가금액, 손익률 다시 표시"],
], widths=[0.7, 1.8, 4.7])
add_code("""// MiniHandler.java - HTTP 요청을 기능별 메서드로 연결
if ("GET".equals(method) && "/api/state".equals(path)) {
    send(exchange, 200, "application/json", project.stateJson());
}
String json = switch (path) {
    case "/api/stock/buy" -> project.buyStock(body);
    case "/api/stock/sell" -> project.sellStock(body);
    default -> Json.obj("ok", false, "error", "없는 API");
};""")
add_code("""// external/KisQuotePoller.java - KIS 현재가 REST API 호출
HttpRequest request = HttpRequest.newBuilder(uri)
    .header("authorization", "Bearer " + token)
    .header("appkey", config.appKey)
    .header("tr_id", "FHKST01010100")
    .GET()
    .build();
HttpResponse<String> response = client.send(request, BodyHandlers.ofString());""")
add_code("""// MiniProject.java - 매수 처리와 거래 기록 저장
Member member = defaultMember();
Stock stock = member.stocks.get(stockName);
long total = (long) stock.price * quantity;
if (member.balance < total) return Json.obj("ok", false);
member.balance -= total;
member.shares.put(stockName, share.buy(quantity, total));
logs.add(new TradeLog(member.uid, stockName, quantity, total, "구매"));
saveDatabase();""")
add_code("""// repository/MySqlDatabase.java - MySQL 트랜잭션 저장
try (Connection con = DriverManager.getConnection(url, user, password)) {
    con.setAutoCommit(false);
    insertMembers(con, members);
    insertShares(con, members);
    insertTradeLogs(con, logs);
    con.commit();
}""")

doc.add_heading("3.4 상속과 인터페이스 사용 이유", level=2)
paragraph("과제 조건상 클래스와 인터페이스 수가 많아야 했기 때문에 단순히 빈 클래스를 늘리기보다 종목 회사 설명과 업종 분류 구조를 분리했다. CompanyProfile은 회사명과 업종처럼 공통 속성을 갖는 추상 클래스이고, StockCategoryProfile은 업종명과 위험 설명을 제공하는 인터페이스다. 이 정보는 MiniProject.addStock()에서 종목 상세 설명을 보강할 때 실제로 사용된다.")
add_table(["구분", "설계", "이유"], [
    ["추상 클래스", "CompanyProfile", "종목별 회사 설명 구조를 통일"],
    ["구현 클래스", "SamsungElectronicsProfile 등", "회사별 설명과 업종 정보를 분리 관리"],
    ["인터페이스", "StockCategoryProfile", "업종별 분류와 위험 설명 제공 규칙을 통일"],
    ["구현 클래스", "업종·테마별 Category 클래스", "종목을 업종과 테마 기준으로 설명할 수 있게 구성"],
], widths=[1.2, 2.5, 3.5])

doc.add_heading("3.5 AI vs 나의 역할", level=2)
add_table(["영역", "내가 한 결정", "AI 활용"], [
    ["주제 방향", "모의주식투자 웹사이트로 확정", "구현 방식 후보 정리"],
    ["데이터 선택", "한국투자증권 KIS Open API 사용 결정", "API 호출 구조와 Java 코드 작성 보조"],
    ["UI 요구", "종목 클릭 중심, 검색/즐겨찾기, 상세 매매 흐름 요구", "HTML/CSS/JS 구현 보조"],
    ["문제 발견", "가격 이상, 불필요 기능, README 문구, 서버 실행 문제 지적", "원인 분석, 코드 수정, 문서 정리"],
], widths=[1.3, 3.0, 3.0])


doc.add_heading("4. 데이터 흐름과 사용자 시나리오", level=1)
paragraph("사용자는 시장 시세 목록에서 거래량이 많은 종목을 확인하고, 검색이나 즐겨찾기로 원하는 종목을 찾는다. 종목을 클릭하면 현재가, 등락률, 거래량, 회사 설명, 가격 변화 그래프가 표시된다. 매수는 종목 상세 화면에서 진행하고, 매도는 보유 탭에서 사용자가 가진 수량을 기준으로 진행한다.")
add_numbered([
    "웹사이트에 접속하면 모의 계좌 화면이 바로 열린다.",
    "시장 시세에서 거래량 상위 종목을 확인하거나 검색한다.",
    "관심 종목을 즐겨찾기에 추가하거나 상세 화면을 연다.",
    "상세 화면에서 가격 그래프를 확인한다.",
    "매수 수량을 입력해 주문한다.",
    "보유 탭에서 평가금액, 손익, 수익률을 확인한다.",
    "보유 종목을 수량 지정 후 매도한다.",
])
add_table(["단계", "내용"], [
    ["입력", "종목 선택, 즐겨찾기, 매수/매도 수량"],
    ["외부 입력", "KIS 현재가, 전일대비, 등락률, 누적 거래량"],
    ["처리", "잔액 확인, 보유 수량 확인, 매매 체결, 손익 계산"],
    ["저장", "모의 계좌 정보, 보유 주식, 거래 기록을 MySQL 우선 + TSV fallback으로 저장"],
    ["출력", "시장 시세, 종목 상세, 그래프, 포트폴리오, 거래 기록"],
], widths=[1.4, 5.8])


doc.add_heading("5. 사용자 UI / 화면", level=1)
paragraph("UI는 실제 투자 앱의 흐름을 참고하되 과제 프로젝트에서 설명하기 쉽도록 단순하게 정리했다. 사용자가 종목을 직접 고르고 판단한 뒤 매수하도록 만들었고, 매도는 보유 종목 기준으로만 진행되게 했다.")
add_bullets([
    "상단 요약: 보유 현금, 주식 평가액, 총자산, 실시간 손익, 수익률 표시",
    "시장 시세: 거래량 상위 종목을 10개 단위 페이지로 표시",
    "검색/즐겨찾기: 종목명, 코드, 업종 검색과 즐겨찾기 목록 분리",
    "종목 상세: 현재가, 등락률, 거래량, 회사 설명, 가격 그래프, 매수 입력",
    "보유 탭: 보유 종목별 평가금액, 손익, 수익률, 매도 수량 입력",
    "기록 탭: 매수/매도 시간, 종목, 수량, 금액 확인",
    "최종 UI: 실제 증권 웹앱을 참고한 흰 배경, 표 중심, 상승 초록/하락 빨강 색상 체계",
])
for image_path, caption_text in UI_SCREENSHOTS:
    if image_path.exists():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run()
        width = Inches(3.4) if "market" in image_path.name else Inches(6.4)
        r.add_picture(str(image_path), width=width)
        caption = doc.add_paragraph(caption_text)
        caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
        caption.runs[0].font.size = Pt(8.5)
        caption.runs[0].font.color.rgb = RGBColor(102, 112, 133)


doc.add_heading("6. 한 달간의 시행착오와 문제 해결", level=1)
add_table(["문제", "해결"], [
    ["실제 시세 연동 문제", "초기에는 KIS REST API가 켜져 있어도 일부 호출 실패 뒤 내장 모의 시세가 같이 실행되어 가격을 덮는 문제가 있었다. 이후 KIS 설정이 있으면 모의 시세로 전환하지 않게 하고, KIS 출처가 찍힌 종목은 모의 Tick이 덮어쓰지 못하게 수정했다."],
    ["전체 종목 처리 범위 결정", "KIS에서 제공하는 전체 종목을 모두 계속 갱신하는 것은 개인 PC에서 실행하는 과제 프로젝트에는 부담이 컸다. 그래서 거래량 상위 종목 중심으로 자동 선별하고, 화면에서는 10개씩 페이지로 나누는 구조를 선택했다."],
    ["UI 흐름 개선", "처음에는 매매 칸에서 종목을 고르는 방식이라 사용자가 불편했다. 이후 시장 시세에서 종목을 클릭하고, 상세 화면에서 가격 그래프를 확인한 뒤 매수하며, 매도는 보유 종목 기준으로 진행하도록 바꿨다."],
    ["뉴스 기능 제거 판단", "네이버 뉴스 API를 붙이는 시도도 했지만 종목마다 기사 품질 편차가 컸고, ETF나 레버리지 상품은 관련 기사 연결이 애매했다. API 키 관리 부담도 있어 최종 버전에서는 모의투자 목적에 맞는 기능만 남기기로 했다."],
    ["DB 저장 전환", "메모리 기반 구조에서는 서버를 재시작하면 모의 계좌, 보유 주식, 거래 기록이 사라졌다. 이를 해결하기 위해 MySQL 테이블 구조를 설계하고 거래 기록과 보유 종목을 저장하는 방식으로 바꿨다."],
], widths=[2.1, 5.1])


doc.add_heading("7. Java의 다양한 클래스 활용", level=1)
add_table(["Java 기능", "사용 내용"], [
    ["HttpServer", "별도 프레임워크 없이 웹 서버 구현"],
    ["HttpClient", "KIS API 호출"],
    ["Thread", "백그라운드 시세 갱신과 소켓 서버 흐름 처리"],
    ["ConcurrentHashMap", "모의 계좌, 종목, 보유 주식 상태 동시 접근 처리"],
    ["ArrayList / List", "거래 기록, 가격 히스토리, 종목 목록 처리"],
    ["LinkedHashMap", "조회 순서와 JSON 응답 순서 유지"],
    ["JDBC DriverManager", "MySQL 연결과 SQL 실행"],
    ["LocalDateTime", "거래 시간, 시세 갱신 시각 표시"],
    ["AtomicLong", "모의 계좌 번호와 시세 Tick 카운트 관리"],
], widths=[2.2, 5.0])


doc.add_heading("8. 데이터 처리", level=1)
paragraph("주식 현재가와 거래량은 KIS API를 기준으로 가져오도록 구현했다. 다만 실제 서비스처럼 모든 종목을 계속 갱신하면 개인 PC에서 실행하는 과제 프로젝트에는 부담이 커진다. 그래서 거래량 상위 종목을 우선 선별하고, 화면에서 필요한 종목 중심으로 가격 정보를 갱신하는 방식이 더 적합하다고 판단했다.")
add_table(["데이터", "처리 방식"], [
    ["주식 시세", "KIS Open API 현재가 REST 조회"],
    ["종목 선별", "거래량 기준 상위 종목 중심으로 자동 갱신"],
    ["실시간성", "전체 종목 매초 조회 대신 상위 목록 중심 갱신과 소켓 구독 구조 실험"],
    ["소켓", "모의 증권사 서버의 가격 Tick 구독 구조"],
    ["사용자 데이터", "MySQL 우선 저장, 연결 실패 시 TSV 파일 저장"],
    ["가격 그래프", "서버가 보관한 최근 가격 히스토리 표시"],
], widths=[2.0, 5.2])

doc.add_heading("8.1 MySQL 테이블 구상", level=2)
add_code("""members(uid, name, balance)
shares(member_uid, stock_name, quantity, purchase_price)
trade_logs(id, member_uid, stock_name, quantity, price, trade_type, traded_at)""")
paragraph("members는 모의 계좌의 현금 잔액을 보관한다. shares는 종목별 보유 수량과 총 매입금액을 저장한다. 평균단가를 별도 컬럼으로 저장하지 않은 이유는 quantity와 purchase_price가 있으면 평균단가를 purchase_price / quantity로 계산할 수 있기 때문이다. trade_logs는 매수와 매도 기록을 시간순으로 남겨 나중에 거래 내역과 수익률 분석을 확장할 수 있게 한다.")
add_table(["테이블", "핵심 컬럼", "설계 이유"], [
    ["members", "uid, name, balance", "단일 모의 계좌의 식별값과 현금 잔액을 유지"],
    ["shares", "member_uid, stock_name, quantity, purchase_price", "보유 수량과 총 매입금액을 저장해 평가금액과 평균단가 계산"],
    ["trade_logs", "id, member_uid, stock_name, quantity, price, trade_type, traded_at", "매수/매도 기록을 남겨 거래 이력과 통계 확장에 활용"],
], widths=[1.3, 3.0, 3.0])

doc.add_heading("8.2 주요 환경변수", level=2)
add_table(["환경변수", "기본값", "역할"], [
    ["KIS_APP_KEY", "없음", "KIS Open API 앱 키"],
    ["KIS_APP_SECRET", "없음", "KIS Open API 앱 시크릿"],
    ["KIS_BASE_URL", "https://openapi.koreainvestment.com:9443", "KIS REST API 주소"],
    ["KIS_MARKET_LIMIT", "200", "거래량 순위 조회 범위, 코드에서 100~300개로 제한"],
    ["KIS_POLL_LIMIT", "100", "REST 폴링 갱신 대상 수, KIS_MARKET_LIMIT 이하로 제한"],
    ["KIS_USE_WEBSOCKET", "false", "true이면 KIS WebSocket 구독을 먼저 시도"],
    ["KIS_WS_URL", "ws://ops.koreainvestment.com:21000", "KIS WebSocket 주소"],
    ["KIS_WS_MAX_SUBSCRIPTIONS", "100", "WebSocket 구독 대상 최대 수"],
    ["MYSQL_URL / MYSQL_USER", "선택", "있으면 MySQL 사용, 없거나 연결 실패 시 TSV fallback"],
    ["MYSQL_PASSWORD", "계정별 설정", "로컬 MySQL 계정에 비밀번호가 있으면 입력, 없으면 생략 가능"],
], widths=[2.2, 2.2, 3.0])

doc.add_heading("8.3 실제 시세 연동 여부 확인", level=2)
paragraph("발표 중 현재 화면의 가격이 실제 API 값인지 질문이 나올 수 있다. 이때는 /api/state 응답의 broker.source와 종목별 quoteSource 값을 보면 된다. KIS 환경변수가 설정된 상태에서는 broker.source가 한국투자증권 KIS REST API로 표시되고, 종목별 quoteSource도 한국투자증권 KIS로 표시된다.")
add_code('''"broker": {
  "source": "한국투자증권 KIS REST API",
  "protocol": "KIS REST API"
},
"quoteSource": "한국투자증권 KIS"''')
add_table(["상황", "동작", "발표 때 설명"], [
    ["KIS 환경변수 있음", "KIS REST API 현재가 사용", "표시 가격은 현재가 API의 stck_prpr 값"],
    ["KIS 환경변수 없음", "내장 모의 증권사 소켓 사용", "API 키 없이도 화면 시연 가능"],
    ["KIS 호출 일부 실패", "REST 폴링 유지", "KIS 출처 종목은 모의 Tick이 덮지 않음"],
], widths=[2.0, 2.5, 2.7])


doc.add_heading("9. 소켓 서버 구조", level=1)
paragraph("프로젝트에는 REST API만 사용하는 구조에서 한 단계 더 나아가기 위해 모의 증권사 소켓 서버 구조와 KIS WebSocket 구독 시도 구조를 추가했다. external 패키지의 MockBrokerServer와 BrokerFeedClient는 ServerSocket과 클라이언트 구독 명령을 이용해 특정 종목의 가격 Tick을 전달한다. KisWebSocketQuoteClient는 승인키 발급, 국내주식 체결 구독 메시지 전송, 체결 메시지 파싱, BrokerTick 변환 흐름을 갖는다.")
add_bullets([
    "서버는 지정 포트에서 소켓 연결을 받고 클라이언트별 요청을 처리한다.",
    "클라이언트는 구독 명령을 보내고 서버는 가격 Tick을 전송한다.",
    "웹 서버는 이 흐름을 받아 포트폴리오 평가와 가격 히스토리에 반영할 수 있다.",
    "KIS WebSocket 시작 또는 실행 중 오류가 발생하면 REST 폴링으로 자동 전환되도록 보완했다.",
    "다만 실제 KIS 테스트베드에서 구독 성공 여부와 메시지 필드 순서는 별도 실계정/모의계정 환경에서 추가 검증이 필요하다.",
])

doc.add_heading("9.1 로그인 기능 제거 판단", level=2)
paragraph("초기에는 회원가입과 로그인 기능을 넣었지만, 이 프로젝트는 실제 사용자에게 배포하는 서비스가 아니라 과제 발표용 모의투자 프로그램이다. 그래서 인증 기능보다 시세 조회, 종목 선택, 매수/매도, 손익 계산 흐름을 설명하는 것이 더 중요하다고 판단했다. 최종 버전에서는 서버가 MySQL에 저장된 모의 계좌 한 명을 사용하고, 사용자는 로그인 없이 바로 모의투자를 진행한다.")
add_bullets([
    "발표 시작 후 바로 시장 시세와 포트폴리오 화면을 보여줄 수 있다.",
    "DB에는 모의 계좌의 현금, 보유 종목, 거래 기록만 저장한다. MySQL을 우선 사용하고, 연결할 수 없으면 TSV 파일 저장소로 전환한다.",
    "회원 인증과 비밀번호 보안 설명보다 Java 서버, KIS API, MySQL 저장 흐름에 집중한다.",
    "실제 배포 서비스로 확장한다면 그때 회원 인증, 세션, HTTPS를 추가하는 것이 적절하다.",
])


doc.add_heading("10. 실제 동작 기능", level=1)
add_table(["기능", "현재 구현 상태"], [
    ["모의 계좌", "로그인 없이 단일 모의 계좌 데이터로 바로 실행"],
    ["시장 시세", "거래량 상위 종목 중심 표시, 10개 단위 페이지 구성"],
    ["검색", "종목명, 코드, 업종 검색 가능"],
    ["즐겨찾기", "브라우저 localStorage 기반 관심 종목 분리 표시"],
    ["종목 상세", "현재가, 등락률, 거래량, 회사 설명, 그래프 표시"],
    ["매수", "선택 종목 상세 화면에서 수량 입력 후 매수"],
    ["매도", "보유 탭에서 보유 수량 기준으로 매도"],
    ["포트폴리오", "현금, 평가금액, 총자산, 손익, 수익률 계산"],
    ["거래 기록", "매수/매도 기록을 MySQL 우선 + TSV fallback으로 저장"],
], widths=[1.9, 5.3])


doc.add_heading("11. 1분 이내 시연 영상", level=1)
paragraph("시연 영상은 실제 코드 설명보다 사용자가 보는 화면 흐름을 보여주는 용도이다. 이번 제출용 영상은 약 35초 길이로 접속, KIS 현재가가 들어온 종목 검색, 가격 그래프 확인, 즐겨찾기, 매수, 보유 탭 확인, 매도, 페이지 이동 장면이 들어 있다.")
add_table(["구간", "보여줄 화면", "설명"], [
    ["0~8초", "웹사이트 접속과 시장 시세", "KIS 현재가와 거래량 기준 종목 목록 확인"],
    ["8~16초", "종목코드 검색", "005930을 입력해 삼성전자 종목을 조회"],
    ["16~25초", "상세 화면과 주문", "현재가, 그래프, 즐겨찾기, 1주 매수"],
    ["25~35초", "보유/기록/매도와 페이지", "보유 탭, 기록 탭, 매도 후 시장 페이지 이동"],
], widths=[1.2, 2.4, 3.6])
paragraph(f"영상 파일 위치: {VIDEO.as_posix() if VIDEO.exists() else 'deliverables/mock-stock-website-demo.webm'}")


doc.add_heading("12. 발표 시 강조할 점", level=1)
paragraph("발표에서는 기능을 많이 넣었다는 설명보다 왜 이 구조가 모의주식투자라는 목적에 맞는지를 먼저 말하는 것이 좋다. 실제 시세와 거래량을 보고 종목을 고른 뒤 매수/매도하는 흐름이 프로젝트의 중심이다.")
add_bullets([
    "실제 증권 API를 붙여 모의투자 화면의 목적을 분명하게 만들었다.",
    "Java 표준 기능만으로 웹 서버, 외부 API, DB 저장, 쓰레드, 소켓 구조를 연결했다.",
    "사용자가 불편하게 느낄 수 있는 드롭다운 주문 방식을 종목 클릭 중심 UI로 바꿨다.",
    "약 2700개 전체 종목을 무리하게 갱신하지 않고 거래량 상위 종목 중심으로 현실적인 범위를 잡았다.",
    "저장, 인코딩, API 키, 가격 표시, 뉴스 기능 제외 판단 같은 실제 개발 문제를 해결했다.",
])


doc.add_heading("13. 마무리와 향후 개선", level=1)
paragraph("최종 결과물은 Java로 직접 만든 모의주식투자 웹앱이다. 한국투자증권 KIS API로 시세와 거래량을 가져오고, 모의 계좌의 보유주식과 거래기록은 MySQL을 우선 사용해 저장한다. MySQL 연결이 불가능한 환경에서는 TSV 파일 저장소로 전환된다. 사용자는 종목을 검색하거나 즐겨찾기하고, 상세 화면에서 가격 그래프를 확인한 뒤 매수/매도를 진행할 수 있다.")
paragraph("향후 개선한다면 실제 KIS 테스트베드에서 WebSocket 구독 성공 여부를 확인하고, service 패키지 안의 MiniProject를 TradingService, PortfolioService, MarketService로 더 세분화할 수 있다. 이후에는 테스트 코드 추가, 회원 인증이 필요한 배포 버전, 업종별 위험 등급 UI 표시, 개인화 관심종목 추천 같은 기능을 확장할 수 있다.")


doc.save(OUT / os.environ.get("REPORT_OUTPUT", "Java_모의주식투자_개인프로젝트_보고서.docx"))

